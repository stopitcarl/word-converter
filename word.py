

import os
import pathlib
import re
from glob import glob
from pathlib import Path

from docx import Document
import win32com.client as win32
from bs4 import BeautifulSoup as bs
from docx.shared import RGBColor
from win32com.client import constants
from tqdm import tqdm


def change_text(run):
    if len(run.text) == 1 and run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
        run.bold = True
        run.underline = True
        run.font.color.rgb = None

def edit_word(filename):    
    # Open the doc
    doc = Document(filename)    

    # Look for red letters in normal text
    for para in doc.paragraphs:
        for run in para.runs:
            # print(run.text)
            change_text(run)

    # Look for red letters in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # print(paragraph.text)
                    for run in paragraph.runs:
                        # print(run.text)
                        change_text(run)

   # Save the file   
    # new_filename = re.sub(r'([a-zA-Z1-9 -_]+).docx$', r'\1-new.docx', filename)
    doc.save(filename)
    

def save_as_docx(path):
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate ()

    # Rename path with .docx    
    new_file_abs = path
    new_file_abs = re.sub(r'\.\w+$', '-automatic.docx', new_file_abs)    

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)
    
    return new_file_abs


def save_as_doc(path):
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate ()

    # Rename path with .doc
    new_file_abs = path
    new_file_abs = re.sub(r'-automatic.docx$', '-new.doc', new_file_abs)
    # print(new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatDocument
    )
    doc.Close(False)
    
    return new_file_abs



# Create list of paths to .doc files
paths = list(Path('.').rglob('*.doc'))
try:
    with tqdm(paths) as t:
        for path in t:
            # Edit description
            t.set_description(f'Editing {path.name}')     
            filename = str(path.absolute())    
            print(filename)    
            docx_file = save_as_docx(filename)
            #print(docx_file)
            edit_word(docx_file)
            doc_file = save_as_doc(docx_file)                    
except Exception as err:
    print("ERROR", err)


# Remove all docx
for path in tqdm(Path('.').rglob('*-automatic.docx'), desc="Removing temporary .docx files"):
    filename = str(path.absolute())
    os.remove(filename)



i = ""
while i != "thanks":
    i = input("Say thanks:\n")


